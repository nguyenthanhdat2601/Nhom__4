import pandas as pd
from numpy import array
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

df=pd.read_csv('diemPython.csv',index_col=0,header = 0)
in_data = array(df.iloc[:,:])

#muc 1: tong so sinh vien
def sum_sv(in_data):
    tong_sv = np.sum(in_data[:, 1])
    return tong_sv

#muc 2: tong so sinh vien dat tung loai diem
def tong_sv_type(in_data):
    Text = "Loại A,Loại B+,Loại B,Loại C+,Loại C,Loại D+,Loại D,Loại F"
    loai_diem = Text.split(",")
    tong_Ldiem = []
    for i in range(3, 11):
        tong_Ldiem.append(np.sum(in_data[:, i]))
    result = dict(zip(loai_diem, tong_Ldiem))
    return result

#muc 3.1: Phần trăm số SV đạt của môn học
def phan_tram_dat(in_data):
    sv_dat = []
    for i in range(1, 9):
        n = np.array(in_data[i, :])
        sv_dat.append(np.sum(n[3:10]))
    phan_tram_sv_dat = np.sum(sv_dat) / sum_sv(in_data)
    return f'{round((phan_tram_sv_dat*100), 2)}%'

#muc 3:Tìm lớp có số SV đạt >=D nhiều nhất/ ít nhất
def Max_Min_pass_class(in_data):
    sv_dat = []
    option = ['Max', 'Min']
    lop2 = []
    for i in range(1, 9):
        n = np.array(in_data[i, :])
        sv_dat.append(np.sum(n[3:10]))
    lop1 = in_data[:, 0]
    a1 = sv_dat.index(max(sv_dat))
    lop2.append(lop1[a1])
    a2 = sv_dat.index(min(sv_dat))
    lop2.append(lop1[a2])
    Dict1 = dict(zip(option, lop2))
    return Dict1

#muc 4: lop co nhieu/it diem A, B, C nhat
def ABC_Max_Min(in_data):
    Text = "Loại A,Loại B+,Loại B,Loại C+,Loại C,Loại D+,Loại D,Loại F"
    option1= Text.split(",")
    option2 = ['Max', 'Min']
    lop1 = in_data[:, 0]
    lop3 = []
    for i in range(3, 11):
        my_dict = {}
        lop2 = []
        A = list(in_data[:, i])
        ind = A.index(max(A))
        lop2.append(lop1[ind])
        ind = A.index(min(A))
        lop2.append(lop1[ind])
        my_dict = dict(zip(option2, lop2))
        lop3.append(my_dict)
    my_dict1 = dict(zip(option1, lop3))
    
    return my_dict1

#muc 5: tim diem nao co sinh vien dat nhieu nhat
def Type_max_min(in_data):
    Text = "Loại A,Loại B+,Loại B,Loại C+,Loại C,Loại D+,Loại D,Loại F"
    loai_diem1 = Text.split(",")
    loai_diem2 = []
    option = ['Max', 'Min']
    tong_Ldiem = []
    for i in range(3, 11):
        tong_Ldiem.append(np.sum(in_data[:, i]))
    ind1 = tong_Ldiem.index(max(tong_Ldiem))
    loai_diem2.append(loai_diem1[ind1])
    ind2 = tong_Ldiem.index(min(tong_Ldiem))
    loai_diem2.append(loai_diem1[ind2])
    result = dict(zip(option, loai_diem2))
    return result

#muc 6: Tìm TBC số sv đạt điểm A,B.. của cả 9 lớp
def tbc_loai_diem(in_data):
    Text = "Loại A,Loại B+,Loại B,Loại C+,Loại C,Loại D+,Loại D,Loại F"
    loai_diem1 = Text.split(",")
    tbc = []
    for i in range(3, 11):
        tbc.append(round(np.mean((in_data[:, i])), 0))
    mydict = dict(zip(loai_diem1, tbc))
    return mydict

#mục 7: tổng số sinh viên đạt bài cuối kỳ
def sv_ck(in_data):
    sv_cuoi_ky = in_data[:, 15]
    return np.sum(sv_cuoi_ky)

#muc 8: so sanh so sinh vien dat chuan L1, L2
def ipL1_l2(in_data):
    sv_L1 = np.sum(in_data[:, 12])
    print("so sinh vien dat chuan dau ra L1: ", sv_L1)
    sv_L2 = np.sum(in_data[:, 13])
    print("so sinh vien dat chuan dau ra L2: ", sv_L2)
    if (sv_L1 > sv_L2) :
        return "=> so sinh vien dat L1 lon hon L2"
    elif (sv_L1 < sv_L2):
        return "=> so sinh vien dat L2 lon hon L1"
    else:
        return "=> so sinh vien dat L2 bang L1"

#muc 9: do thi pho diem tung lop
def draw_graph_mark(in_data):
    Ma_lop = list(in_data[:, 0])
    Pho_diem = []
    Pd_lop = {}
    Text = "Loại A,Loại B+,Loại B,Loại C+,Loại C,Loại D+,Loại D,Loại F"
    loai_diem1 = Text.split(",")
    for i in range(0, 9):
        info = list(in_data[i, :])
        type_diem = []
        for j in range(3, 11):
            type_diem.append(info[j])
        my_dict1 = dict(zip(loai_diem1, type_diem))
        Pho_diem.append(my_dict1)
    Pd_lop = dict(zip(Ma_lop, Pho_diem))
    
    data = Pd_lop
    tenlop = list(data.keys())
    Pho_diem = list(data[tenlop[0]].keys())
    x = np.arange(len(tenlop))
    bottom_values = [0] * len(tenlop)
    for loai_diem in Pho_diem:
        values = [data[lop][loai_diem] for lop in tenlop]
        plt.bar(x, values, label=loai_diem, bottom=bottom_values)
        bottom_values = [bottom + value for bottom, value in zip(bottom_values, values)]
    plt.xlabel('Lớp')
    plt.ylabel('số lượng điểm')
    plt.title('Phổ điểm của từng lớp')
    plt.xticks(x, tenlop)
    plt.yticks(np.arange(0, 65, 2))
    plt.legend()
    fig = plt.gcf()
    fig.set_size_inches(17, 8) 
    plt.show()
    return 'Đã hoàn thành'

# muc 10: số phần trăm từng loại điểm và hiển thị bằng biểu đồ hình tròn:
def phantram_loai_diem(in_data):
    diem_loai = df.loc[:, 'Loại A':'Loại F']
    tong_so_sv = diem_loai.sum() # tính tổng số sinh viên
    labels = tong_so_sv.index  # Tên loại điểm (A, B+, B, C+, C, D+, D, F)
    sizes = tong_so_sv.values  # Số lượng sinh viên đạt từng loại điểm
    colors = ['red', 'yellowgreen', 'lightblue', 'darkgreen', 'blue', 'pink', 'lightsalmon', 'brown']
    explode = (0.1, 0, 0, 0, 0, 0, 0, 0)  # Để phân cách một phần nhỏ (loại A) ra xa

    plt.figure(figsize=(8, 8))
    plt.pie(sizes, explode=explode, labels=labels, colors=colors, autopct='%1.1f%%', shadow=True, startangle=140)
    plt.title('Phan phoi diem theo loai')
    plt.axis('equal')  # Đảm bảo biểu đồ hình tròn hoàn chỉnh
    plt.show()
    return 'Đã hoàn thành'
# muc 11: Tính số phần trăm sinh viên đi thi đầy đủ
def phantram_sv_tham_gia_day_du(in_data):
    so_sv_du = np.sum(in_data[:, 15])
    so_sv_khong_du = np.sum(in_data[:, 1]) - so_sv_du

    # Tạo biểu đồ hình tròn
    labels = ['Tham gia đầy đủ', 'Tham gia không đầy đủ']
    sizes = [so_sv_du, so_sv_khong_du]
    colors = ['lightgreen', 'lightcoral']
    explode = (0.1, 0)  # Để phân cách một phần nhỏ (loại Đủ) ra xa     
    plt.figure(figsize=(8, 8))
    plt.pie(sizes, explode=explode, labels=labels, colors=colors, autopct='%1.1f%%', shadow=True, startangle=140)
    plt.title('Phân phối số lượng sinh viên tham gia thi KTHP')
    plt.axis('equal')  # Đảm bảo biểu đồ hình tròn hoàn chỉnh
    plt.show()
    return 'Đã hoàn thành'

#muc 13: xuất báo cáo theo từng mục
def xuat_report(in_data):

    doc = Document()

    title = doc.add_heading('BÁO CÁO KẾT QUẢ KẾT THÚC HỌC PHẦN ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Căn giữa tiêu đề

    # Thông tin học sinh
    teacher_name = 'Ths. Phạm Thị Quỳnh Trang'
    subject = 'Lập trình Python'
    faculty = 'Khoa Điện Tử'

    # Thêm thông tin học sinh vào tài liệu
    doc.add_paragraph(f'Giảng viên: {teacher_name}', style='Heading 1')
    doc.add_paragraph(f'Học phần: {subject}', style='Heading 2')
    doc.add_paragraph(f'Khoa: {faculty}', style='Heading 2')

    # Thêm kết quả học tập
    doc.add_heading('Kết quả học tập', level=1)

    # Thêm bảng điểm
    table = doc.add_table(rows=1, cols=3)
    table.autofit = True
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Đặt tiêu đề cho các cột
    table.rows[0].cells[0].text = 'Môn học'
    table.rows[0].cells[1].text = 'Điểm'
    table.rows[0].cells[2].text = 'Xếp loại'

    # Giả định danh sách môn học, điểm, và xếp loại
    subjects = ['Toán', 'Văn', 'Anh Văn', 'Lý', 'Hóa']
    grades = [8.5, 7.75, 9.0, 8.0, 7.5]

    # Thêm thông tin môn học, điểm, và xếp loại vào bảng
    for subject, grade in zip(subjects, grades):
        row = table.add_row().cells
        row[0].text = subject
        row[1].text = str(grade)
        if grade >= 8.0:
            row[2].text = 'Giỏi'
        elif grade >= 6.5:
            row[2].text = 'Khá'
        else:
            row[2].text = 'Trung bình'

    # Lưu tài liệu vào tệp Word
    doc.save('bao_cao_ket_qua_hoc_tap.docx')

    print('Báo cáo đã được tạo và lưu thành công.')
